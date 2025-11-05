#!/usr/bin/env python3
"""
Apply PacktPub [PACKT] Styles to Pandoc-Converted Word Documents v3

FIXED: Properly detects and converts list items (with numbering properties)

PacktPub Style System:
- Headings use standard "Heading 1-6" (NO [PACKT] suffix)
- All other content uses [PACKT] suffix (Normal [PACKT], Code [PACKT], etc.)
- Lists have numbering properties but Pandoc assigns "Normal" style

Usage:
    python3 apply-packt-styles-v3.py input.docx output.docx
"""

from docx import Document
from docx.oxml import parse_xml
import sys
from pathlib import Path

# Mapping from Pandoc's built-in styles to PacktPub styles
STYLE_MAPPINGS = {
    # Paragraph styles (add [PACKT] suffix)
    'Normal': 'Normal [PACKT]',
    'Source Code': 'Code [PACKT]',
    'Code': 'Code [PACKT]',
    'Block Quote': 'Quote [PACKT]',
    'Quote': 'Quote [PACKT]',

    # List styles (add [PACKT] suffix)
    # Note: Pandoc creates lists with numbering but keeps "Normal" style
    # We detect these by checking for numPr (numbering properties)
    'List Bullet': 'Bullet [PACKT]',
    'List Number': 'Numbered Bullet [PACKT]',
    'List Paragraph': 'Bullet [PACKT]',
}

# Character style mappings (for inline formatting)
CHARACTER_STYLE_MAPPINGS = {
    'Strong': 'Key Word [PACKT]',
    'Emphasis': 'Italics [PACKT]',
    'Verbatim Char': 'Code In Text [PACKT]',
    'Source Text': 'Code In Text [PACKT]',
}

def has_numbering(paragraph):
    """Check if paragraph has numbering properties (is part of a list)"""
    if paragraph._element.pPr is None:
        return False

    numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    return numPr is not None

def get_numbering_level(paragraph):
    """Get the numbering level (0-based) of a list item"""
    if not has_numbering(paragraph):
        return None

    numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    ilvl = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')

    if ilvl is not None:
        return int(ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0))

    return 0

def is_numbered_list(paragraph):
    """Check if this is a numbered list (vs bullet list)"""
    if not has_numbering(paragraph):
        return False

    # In Pandoc output, numbered lists typically have numId
    # This is a heuristic - may need refinement
    numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    numId = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')

    if numId is not None:
        num_id_val = int(numId.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0))
        # Numbered lists typically have odd numIds in Pandoc output (heuristic)
        return num_id_val % 2 == 1

    return False

def apply_packt_styles(input_file, output_file):
    """
    Apply PacktPub [PACKT] styles to a Word document.

    Args:
        input_file: Path to input .docx file (Pandoc output)
        output_file: Path to output .docx file (with PACKT styles)
    """
    print(f"Loading document: {input_file}")
    doc = Document(input_file)

    # Check if PACKT styles exist in the document
    available_styles = {style.name for style in doc.styles}
    packt_styles = {s for s in available_styles if '[PACKT]' in s}

    if not packt_styles:
        print("⚠️  WARNING: No [PACKT] styles found in document!")
        print("   Make sure you used Sample Chapter.docx as --reference-doc")
        return False

    print(f"✓ Found {len(packt_styles)} [PACKT] styles in document")

    # Apply paragraph style mappings
    para_count = 0
    para_mapped = 0
    para_skipped_headings = 0
    lists_mapped = 0

    for para in doc.paragraphs:
        para_count += 1
        current_style = para.style.name

        # Skip headings - PacktPub uses standard "Heading 1-6" without [PACKT]
        if current_style.startswith('Heading'):
            para_skipped_headings += 1
            continue

        # Check if this is a list item (has numbering)
        if has_numbering(para):
            # Determine if bullet or numbered list
            if is_numbered_list(para):
                target_style = 'Numbered Bullet [PACKT]'
            else:
                target_style = 'Bullet [PACKT]'

            if target_style in available_styles:
                para.style = target_style
                lists_mapped += 1
            else:
                print(f"⚠️  List style not found: {target_style}")
            continue

        # Regular paragraph style mapping
        if current_style in STYLE_MAPPINGS:
            target_style = STYLE_MAPPINGS[current_style]

            # Check if target style exists
            if target_style in available_styles:
                para.style = target_style
                para_mapped += 1
            else:
                print(f"⚠️  Style not found: {target_style}")

    print(f"\n✓ Processed {para_count} paragraphs")
    print(f"✓ Mapped {para_mapped} paragraph styles to [PACKT] equivalents")
    print(f"✓ Mapped {lists_mapped} list items to Bullet/Numbered Bullet [PACKT]")
    print(f"✓ Kept {para_skipped_headings} headings as standard \"Heading X\" (correct for PacktPub)")

    # Apply character style mappings (runs within paragraphs)
    run_count = 0
    run_mapped = 0

    for para in doc.paragraphs:
        for run in para.runs:
            run_count += 1
            if run.style and run.style.name in CHARACTER_STYLE_MAPPINGS:
                target_style = CHARACTER_STYLE_MAPPINGS[run.style.name]
                if target_style in available_styles:
                    run.style = target_style
                    run_mapped += 1

    if run_mapped > 0:
        print(f"✓ Mapped {run_mapped} / {run_count} character styles to [PACKT] equivalents")

    # Save output
    print(f"\nSaving to: {output_file}")
    doc.save(output_file)
    print("✅ Complete!")

    return True

def print_style_report(input_file):
    """Print a report of styles used in the document"""
    doc = Document(input_file)

    style_usage = {}
    for para in doc.paragraphs:
        style_name = para.style.name
        style_usage[style_name] = style_usage.get(style_name, 0) + 1

    print("\n" + "=" * 70)
    print("STYLE USAGE REPORT")
    print("=" * 70)

    for style, count in sorted(style_usage.items(), key=lambda x: -x[1]):
        if style.startswith('Heading'):
            marker = ' ✓ (PacktPub standard)'
        elif '[PACKT]' in style:
            marker = ' ✓ [PACKT]'
        else:
            marker = ' ⚠️  (needs mapping)'
        print(f"  {count:3}x  {style:45} {marker}")

    packt_count = sum(1 for s in style_usage if '[PACKT]' in s)
    heading_count = sum(1 for s in style_usage if s.startswith('Heading'))
    total_count = len(style_usage)
    correct_count = packt_count + heading_count

    print(f"\nPacktPub-ready styles: {correct_count} / {total_count} style types")
    print(f"  • [PACKT] styles: {packt_count}")
    print(f"  • Standard headings: {heading_count}")
    print("=" * 70 + "\n")

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 apply-packt-styles-v3.py input.docx output.docx")
        print("\nApplies PacktPub [PACKT] styles to Pandoc-converted documents")
        print("\nPacktPub Style System:")
        print("  • Headings: Standard 'Heading 1-6' (no [PACKT])")
        print("  • Lists: Bullet [PACKT] / Numbered Bullet [PACKT]")
        print("  • Other content: Add [PACKT] suffix")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2])

    if not input_file.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    print("PacktPub Style Application Tool v3")
    print("=" * 70 + "\n")

    # Apply styles
    success = apply_packt_styles(input_file, output_file)

    if success:
        # Print report
        print_style_report(output_file)
    else:
        sys.exit(1)

if __name__ == "__main__":
    main()
