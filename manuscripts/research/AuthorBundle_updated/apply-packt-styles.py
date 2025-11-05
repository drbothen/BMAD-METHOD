#!/usr/bin/env python3
"""
Apply PacktPub [PACKT] Styles to Pandoc-Converted Word Documents

This script post-processes Word documents converted by Pandoc to apply
PacktPub's [PACKT] styles instead of the built-in Word styles that Pandoc uses.

Usage:
    python3 apply-packt-styles.py input.docx output.docx
"""

from docx import Document
import sys
from pathlib import Path

# Mapping from Pandoc's built-in styles to PacktPub [PACKT] styles
STYLE_MAPPINGS = {
    # Heading styles
    'Heading 1': 'Heading 1 [PACKT]',
    'Heading 2': 'Heading 2 [PACKT]',
    'Heading 3': 'Heading 3 [PACKT]',
    'Heading 4': 'Heading 4 [PACKT]',
    'Heading 5': 'Heading 5 [PACKT]',
    'Heading 6': 'Heading 6 [PACKT]',

    # Paragraph styles
    'Normal': 'Normal [PACKT]',
    'Source Code': 'Code [PACKT]',
    'Code': 'Code [PACKT]',
    'Block Quote': 'Quote [PACKT]',
    'Quote': 'Quote [PACKT]',

    # List styles
    'List Bullet': 'Bullet [PACKT]',
    'List Number': 'Numbered Bullet [PACKT]',
    'List Paragraph': 'Bullet [PACKT]',  # Pandoc sometimes uses this
}

# Character style mappings (for inline formatting)
CHARACTER_STYLE_MAPPINGS = {
    'Strong': 'Key Word [PACKT]',
    'Emphasis': 'Italics [PACKT]',
    'Verbatim Char': 'Code In Text [PACKT]',
    'Source Text': 'Code In Text [PACKT]',
}

def apply_packt_styles(input_file, output_file):
    """
    Apply PacktPub [PACKT] styles to a Word document.

    Args:
        input_file: Path to input .docx file (Pandoc output)
        output_file: Path to output .docx file (with PACKT styles)
    """
    print(f"Loading document: {input_file}")
    doc = Document(input_file)

    # Check if PACKT styles exist in the document
    available_styles = {style.name for style in doc.styles}
    packt_styles = {s for s in available_styles if '[PACKT]' in s}

    if not packt_styles:
        print("⚠️  WARNING: No [PACKT] styles found in document!")
        print("   Make sure you used Sample Chapter.docx as --reference-doc")
        return False

    print(f"✓ Found {len(packt_styles)} [PACKT] styles in document")

    # Apply paragraph style mappings
    para_count = 0
    para_mapped = 0

    for para in doc.paragraphs:
        para_count += 1
        current_style = para.style.name

        if current_style in STYLE_MAPPINGS:
            target_style = STYLE_MAPPINGS[current_style]

            # Check if target style exists
            if target_style in available_styles:
                para.style = target_style
                para_mapped += 1
            else:
                print(f"⚠️  Style not found: {target_style}")

    print(f"\n✓ Processed {para_count} paragraphs")
    print(f"✓ Mapped {para_mapped} paragraph styles to [PACKT] equivalents")

    # Apply character style mappings (runs within paragraphs)
    run_count = 0
    run_mapped = 0

    for para in doc.paragraphs:
        for run in para.runs:
            run_count += 1
            if run.style and run.style.name in CHARACTER_STYLE_MAPPINGS:
                target_style = CHARACTER_STYLE_MAPPINGS[run.style.name]
                if target_style in available_styles:
                    run.style = target_style
                    run_mapped += 1

    if run_mapped > 0:
        print(f"✓ Mapped {run_mapped} / {run_count} character styles to [PACKT] equivalents")

    # Save output
    print(f"\nSaving to: {output_file}")
    doc.save(output_file)
    print("✅ Complete!")

    return True

def print_style_report(input_file):
    """Print a report of styles used in the document"""
    doc = Document(input_file)

    style_usage = {}
    for para in doc.paragraphs:
        style_name = para.style.name
        style_usage[style_name] = style_usage.get(style_name, 0) + 1

    print("\n" + "=" * 70)
    print("STYLE USAGE REPORT")
    print("=" * 70)

    for style, count in sorted(style_usage.items(), key=lambda x: -x[1]):
        packt_marker = ' ✓ [PACKT]' if '[PACKT]' in style else ''
        print(f"  {count:3}x  {style:45} {packt_marker}")

    packt_count = sum(1 for s in style_usage if '[PACKT]' in s)
    total_count = len(style_usage)
    print(f"\nPACKT styles: {packt_count} / {total_count} style types")
    print("=" * 70 + "\n")

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 apply-packt-styles.py input.docx output.docx")
        print("\nApplies PacktPub [PACKT] styles to Pandoc-converted documents")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2])

    if not input_file.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    print("PacktPub Style Application Tool")
    print("=" * 70 + "\n")

    # Apply styles
    success = apply_packt_styles(input_file, output_file)

    if success:
        # Print report
        print_style_report(output_file)
    else:
        sys.exit(1)

if __name__ == "__main__":
    main()
