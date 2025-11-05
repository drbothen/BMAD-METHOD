#!/usr/bin/env python3
"""
Apply PacktPub [PACKT] Styles to Pandoc-Converted Word Documents

PacktPub Style System:
- Headings use standard "Heading 1-6" (NO [PACKT] suffix)
- All other content uses [PACKT] suffix (Normal [PACKT], Code [PACKT], etc.)

Usage:
    python3 apply-packt-styles-v2.py input.docx output.docx
"""

from docx import Document
import sys
from pathlib import Path

# Mapping from Pandoc's built-in styles to PacktPub styles
# NOTE: Headings stay as "Heading 1-6" - PacktPub doesn't use [PACKT] for headings!
STYLE_MAPPINGS = {
    # Paragraph styles (add [PACKT] suffix)
    'Normal': 'Normal [PACKT]',
    'Source Code': 'Code [PACKT]',
    'Code': 'Code [PACKT]',
    'Block Quote': 'Quote [PACKT]',
    'Quote': 'Quote [PACKT]',

    # List styles (add [PACKT] suffix)
    'List Bullet': 'Bullet [PACKT]',
    'List Number': 'Numbered Bullet [PACKT]',
    'List Paragraph': 'Bullet [PACKT]',

    # Headings stay as-is (PacktPub uses standard Heading 1-6)
    # These don't need mapping - Pandoc already outputs correct style names
}

# Character style mappings (for inline formatting)
CHARACTER_STYLE_MAPPINGS = {
    'Strong': 'Key Word [PACKT]',
    'Emphasis': 'Italics [PACKT]',
    'Verbatim Char': 'Code In Text [PACKT]',
    'Source Text': 'Code In Text [PACKT]',
}

def apply_packt_styles(input_file, output_file):
    """
    Apply PacktPub [PACKT] styles to a Word document.

    Args:
        input_file: Path to input .docx file (Pandoc output)
        output_file: Path to output .docx file (with PACKT styles)
    """
    print(f"Loading document: {input_file}")
    doc = Document(input_file)

    # Check if PACKT styles exist in the document
    available_styles = {style.name for style in doc.styles}
    packt_styles = {s for s in available_styles if '[PACKT]' in s}

    if not packt_styles:
        print("⚠️  WARNING: No [PACKT] styles found in document!")
        print("   Make sure you used Sample Chapter.docx as --reference-doc")
        return False

    print(f"✓ Found {len(packt_styles)} [PACKT] styles in document")

    # Apply paragraph style mappings
    para_count = 0
    para_mapped = 0
    para_skipped_headings = 0

    for para in doc.paragraphs:
        para_count += 1
        current_style = para.style.name

        # Skip headings - PacktPub uses standard "Heading 1-6" without [PACKT]
        if current_style.startswith('Heading'):
            para_skipped_headings += 1
            continue

        if current_style in STYLE_MAPPINGS:
            target_style = STYLE_MAPPINGS[current_style]

            # Check if target style exists
            if target_style in available_styles:
                para.style = target_style
                para_mapped += 1
            else:
                print(f"⚠️  Style not found: {target_style}")

    print(f"\n✓ Processed {para_count} paragraphs")
    print(f"✓ Mapped {para_mapped} paragraph styles to [PACKT] equivalents")
    print(f"✓ Kept {para_skipped_headings} headings as standard \"Heading X\" (correct for PacktPub)")

    # Apply character style mappings (runs within paragraphs)
    run_count = 0
    run_mapped = 0

    for para in doc.paragraphs:
        for run in para.runs:
            run_count += 1
            if run.style and run.style.name in CHARACTER_STYLE_MAPPINGS:
                target_style = CHARACTER_STYLE_MAPPINGS[run.style.name]
                if target_style in available_styles:
                    run.style = target_style
                    run_mapped += 1

    if run_mapped > 0:
        print(f"✓ Mapped {run_mapped} / {run_count} character styles to [PACKT] equivalents")

    # Save output
    print(f"\nSaving to: {output_file}")
    doc.save(output_file)
    print("✅ Complete!")

    return True

def print_style_report(input_file):
    """Print a report of styles used in the document"""
    doc = Document(input_file)

    style_usage = {}
    for para in doc.paragraphs:
        style_name = para.style.name
        style_usage[style_name] = style_usage.get(style_name, 0) + 1

    print("\n" + "=" * 70)
    print("STYLE USAGE REPORT")
    print("=" * 70)

    for style, count in sorted(style_usage.items(), key=lambda x: -x[1]):
        if style.startswith('Heading'):
            marker = ' ✓ (PacktPub standard)'
        elif '[PACKT]' in style:
            marker = ' ✓ [PACKT]'
        else:
            marker = ' ⚠️  (needs mapping)'
        print(f"  {count:3}x  {style:45} {marker}")

    packt_count = sum(1 for s in style_usage if '[PACKT]' in s)
    heading_count = sum(1 for s in style_usage if s.startswith('Heading'))
    total_count = len(style_usage)
    correct_count = packt_count + heading_count

    print(f"\nPacktPub-ready styles: {correct_count} / {total_count} style types")
    print(f"  • [PACKT] styles: {packt_count}")
    print(f"  • Standard headings: {heading_count}")
    print("=" * 70 + "\n")

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 apply-packt-styles-v2.py input.docx output.docx")
        print("\nApplies PacktPub [PACKT] styles to Pandoc-converted documents")
        print("\nPacktPub Style System:")
        print("  • Headings: Standard 'Heading 1-6' (no [PACKT])")
        print("  • Other content: Add [PACKT] suffix")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2])

    if not input_file.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    print("PacktPub Style Application Tool v2")
    print("=" * 70 + "\n")

    # Apply styles
    success = apply_packt_styles(input_file, output_file)

    if success:
        # Print report
        print_style_report(output_file)
    else:
        sys.exit(1)

if __name__ == "__main__":
    main()
