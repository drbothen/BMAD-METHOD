#!/usr/bin/env python3
"""
Verify PacktPub Word Document Formatting

Post-conversion validation checks:
- [PACKT] style application correctness
- Code block structure (Code [PACKT] / Code End [PACKT])
- List style application
- Image embedding
- Code block line counts in Word document

Usage:
    python3 verify-packt-document.py output.docx [--report verification-report.md]
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)


class ValidationResult:
    """Stores validation results with severity levels"""

    def __init__(self):
        self.errors = []      # Critical issues (must fix)
        self.warnings = []    # Should fix
        self.info = []        # Nice to have

    def add_error(self, message):
        self.errors.append(message)

    def add_warning(self, message):
        self.warnings.append(message)

    def add_info(self, message):
        self.info.append(message)

    def has_errors(self):
        return len(self.errors) > 0

    def summary(self):
        if len(self.errors) > 0:
            return f"🔴 FAIL - {len(self.errors)} critical issues"
        elif len(self.warnings) > 0:
            return f"🟡 WARNINGS - {len(self.warnings)} warnings"
        else:
            return f"✅ PASS - All checks passed"


def verify_packt_styles(doc):
    """Verify all paragraphs use PacktPub-compliant styles"""
    results = ValidationResult()

    style_usage = {}
    unmapped_paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name
        style_usage[style_name] = style_usage.get(style_name, 0) + 1

        # Check for unmapped styles (neither [PACKT] nor standard Heading)
        if not style_name.startswith('Heading') and '[PACKT]' not in style_name:
            unmapped_paragraphs.append((i + 1, style_name, para.text[:50]))

    if unmapped_paragraphs:
        results.add_error(f"Found {len(unmapped_paragraphs)} paragraphs with unmapped styles:")
        for para_num, style, text in unmapped_paragraphs[:5]:  # Show first 5
            results.add_error(f"  Para {para_num}: '{style}' - {text}...")
        if len(unmapped_paragraphs) > 5:
            results.add_error(f"  ... and {len(unmapped_paragraphs) - 5} more unmapped paragraphs")

    # Count style types
    packt_count = sum(1 for s in style_usage if '[PACKT]' in s)
    heading_count = sum(1 for s in style_usage if s.startswith('Heading'))
    total_count = len(style_usage)
    correct_count = packt_count + heading_count

    results.add_info(f"Style coverage: {correct_count}/{total_count} style types are PacktPub-compliant")
    results.add_info(f"  • [PACKT] styles: {packt_count}")
    results.add_info(f"  • Standard headings: {heading_count}")

    # List most common styles
    top_styles = sorted(style_usage.items(), key=lambda x: -x[1])[:10]
    results.add_info(f"Top styles: " + ", ".join([f"{s}({c})" for s, c in top_styles[:5]]))

    return results


def verify_code_blocks(doc):
    """Verify code block formatting and line counts"""
    results = ValidationResult()

    code_blocks = []
    current_block = []

    # Find code blocks (sequences of Code [PACKT] / Code End [PACKT])
    for i, para in enumerate(doc.paragraphs):
        if para.style.name in ['Code [PACKT]', 'Code End [PACKT]']:
            current_block.append((i + 1, para.style.name, para.text))

            # End of block if Code End [PACKT]
            if para.style.name == 'Code End [PACKT]':
                code_blocks.append(current_block)
                current_block = []
        else:
            # If we were in a block but didn't end with Code End, that's an error
            if current_block:
                results.add_error(f"Code block starting at para {current_block[0][0]} missing Code End [PACKT]")
                current_block = []

    # Validate each block
    for block_num, block in enumerate(code_blocks, 1):
        line_count = len(block)

        # Check last line has Code End [PACKT]
        if block[-1][1] != 'Code End [PACKT]':
            results.add_error(f"Code block #{block_num}: Last line should use 'Code End [PACKT]' but uses '{block[-1][1]}'")

        # Check line count
        if line_count > 30:
            results.add_error(f"Code block #{block_num}: {line_count} lines (MAX: 30)")
        elif line_count > 20:
            results.add_warning(f"Code block #{block_num}: {line_count} lines (IDEAL: ≤20)")

    # Check for single-line "code blocks" that should just be Code [PACKT]
    single_code_end = sum(1 for para in doc.paragraphs if para.style.name == 'Code End [PACKT]')
    if single_code_end != len(code_blocks):
        results.add_warning(f"Code End [PACKT] count ({single_code_end}) != detected code blocks ({len(code_blocks)})")

    results.add_info(f"Total code blocks: {len(code_blocks)}")
    if code_blocks:
        avg_lines = sum(len(block) for block in code_blocks) / len(code_blocks)
        results.add_info(f"Average code block length: {avg_lines:.1f} lines")

    return results


def verify_list_styles(doc):
    """Verify list formatting"""
    results = ValidationResult()

    bullet_count = sum(1 for para in doc.paragraphs if para.style.name == 'Bullet [PACKT]')
    numbered_count = sum(1 for para in doc.paragraphs if para.style.name == 'Numbered Bullet [PACKT]')

    # Check for unmapped lists (still have Normal style but have numbering)
    unmapped_lists = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Normal':
            # Check if has numbering properties
            if para._element.pPr is not None:
                numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    unmapped_lists.append((i + 1, para.text[:50]))

    if unmapped_lists:
        results.add_error(f"Found {len(unmapped_lists)} list items still using 'Normal' style:")
        for para_num, text in unmapped_lists[:5]:
            results.add_error(f"  Para {para_num}: {text}...")
        if len(unmapped_lists) > 5:
            results.add_error(f"  ... and {len(unmapped_lists) - 5} more unmapped lists")

    results.add_info(f"Bullet lists: {bullet_count}")
    results.add_info(f"Numbered lists: {numbered_count}")
    results.add_info(f"Total list items: {bullet_count + numbered_count}")

    return results


def verify_images(doc):
    """Verify image embedding"""
    results = ValidationResult()

    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    if image_count == 0:
        results.add_info("No images embedded in document")
    else:
        results.add_info(f"Total embedded images: {image_count}")

    return results


def verify_heading_structure(doc):
    """Verify heading hierarchy"""
    results = ValidationResult()

    headings = [(i + 1, para.style.name, para.text) for i, para in enumerate(doc.paragraphs)
                if para.style.name.startswith('Heading')]

    if not headings:
        results.add_warning("No headings found in document")
        return results

    # Check for H1 (should be one - the chapter title)
    h1_count = sum(1 for _, style, _ in headings if style == 'Heading 1')
    if h1_count == 0:
        results.add_error("Missing Heading 1 (chapter title)")
    elif h1_count > 1:
        results.add_warning(f"Multiple Heading 1 found ({h1_count}). Should be one chapter title.")

    # Check for level skipping (e.g., H2 -> H4)
    prev_level = 0
    for para_num, style, text in headings:
        level = int(style.split()[1]) if len(style.split()) > 1 else 1

        if level - prev_level > 1:
            results.add_warning(f"Heading level skip at para {para_num}: {style} after Heading {prev_level}")

        prev_level = level

    results.add_info(f"Total headings: {len(headings)}")
    heading_dist = {}
    for _, style, _ in headings:
        heading_dist[style] = heading_dist.get(style, 0) + 1

    results.add_info(f"Heading distribution: " + ", ".join([f"{s}({c})" for s, c in sorted(heading_dist.items())]))

    return results


def generate_report(doc_path, style_results, code_results, list_results, image_results, heading_results, output_path):
    """Generate verification report"""

    all_results = [style_results, code_results, list_results, image_results, heading_results]

    total_errors = sum(len(r.errors) for r in all_results)
    total_warnings = sum(len(r.warnings) for r in all_results)
    total_info = sum(len(r.info) for r in all_results)

    # Determine overall status
    if total_errors > 0:
        overall = "🔴 FAIL"
    elif total_warnings > 0:
        overall = "🟡 WARNINGS"
    else:
        overall = "✅ PASS"

    report = f"""# PacktPub Document Verification Report

**Document**: {doc_path}
**Date**: {Path(doc_path).stat().st_mtime}

## Overall Status

{overall} - {total_errors} errors, {total_warnings} warnings

## Summary

### Critical Issues (MUST FIX): {total_errors}
### Warnings (SHOULD FIX): {total_warnings}
### Information: {total_info}

---

## Style Verification

{style_results.summary()}

"""

    if style_results.errors:
        report += "\n### ❌ Errors\n"
        for error in style_results.errors:
            report += f"- {error}\n"

    if style_results.warnings:
        report += "\n### ⚠️ Warnings\n"
        for warning in style_results.warnings:
            report += f"- {warning}\n"

    if style_results.info:
        report += "\n### ℹ️ Information\n"
        for info in style_results.info:
            report += f"- {info}\n"

    report += "\n---\n\n## Code Block Verification\n\n"
    report += f"{code_results.summary()}\n"

    if code_results.errors:
        report += "\n### ❌ Errors\n"
        for error in code_results.errors:
            report += f"- {error}\n"

    if code_results.warnings:
        report += "\n### ⚠️ Warnings\n"
        for warning in code_results.warnings:
            report += f"- {warning}\n"

    if code_results.info:
        report += "\n### ℹ️ Information\n"
        for info in code_results.info:
            report += f"- {info}\n"

    report += "\n---\n\n## List Style Verification\n\n"
    report += f"{list_results.summary()}\n"

    if list_results.errors:
        report += "\n### ❌ Errors\n"
        for error in list_results.errors:
            report += f"- {error}\n"

    if list_results.warnings:
        report += "\n### ⚠️ Warnings\n"
        for warning in list_results.warnings:
            report += f"- {warning}\n"

    if list_results.info:
        report += "\n### ℹ️ Information\n"
        for info in list_results.info:
            report += f"- {info}\n"

    report += "\n---\n\n## Image Verification\n\n"
    report += f"{image_results.summary()}\n"

    if image_results.info:
        report += "\n### ℹ️ Information\n"
        for info in image_results.info:
            report += f"- {info}\n"

    report += "\n---\n\n## Heading Structure\n\n"
    report += f"{heading_results.summary()}\n"

    if heading_results.errors:
        report += "\n### ❌ Errors\n"
        for error in heading_results.errors:
            report += f"- {error}\n"

    if heading_results.warnings:
        report += "\n### ⚠️ Warnings\n"
        for warning in heading_results.warnings:
            report += f"- {warning}\n"

    if heading_results.info:
        report += "\n### ℹ️ Information\n"
        for info in heading_results.info:
            report += f"- {info}\n"

    report += "\n---\n\n## Next Steps\n\n"

    if total_errors > 0:
        report += "**CRITICAL**: Document has errors that must be fixed before submission.\n\n"
        report += "1. Review errors listed above\n"
        report += "2. Re-run conversion: `./format-for-packtpub.sh <manuscript.md>`\n"
        report += "3. Verify again: `python3 verify-packt-document.py " + str(doc_path) + "`\n"
    elif total_warnings > 0:
        report += "**WARNINGS**: Document has minor issues. Review and fix if possible.\n\n"
        report += "1. Review warnings listed above\n"
        report += "2. Manual fixes in Word if needed\n"
        report += "3. Run PacktPub submission checklist\n"
    else:
        report += "✅ Document is ready for PacktPub submission!\n\n"
        report += "Next steps:\n"
        report += "1. Manual review in Word for special formatting (info boxes, tips, warnings)\n"
        report += "2. Run PacktPub submission checklist\n"
        report += "3. Submit via PacktPub AuthorSight portal\n"

    # Write report
    if output_path:
        with open(output_path, 'w') as f:
            f.write(report)
        print(f"\n✓ Report saved to: {output_path}")
    else:
        print(report)

    return total_errors == 0


def main():
    parser = argparse.ArgumentParser(
        description='Verify PacktPub Word document formatting',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 verify-packt-document.py chapter-05-formatted.docx
  python3 verify-packt-document.py chapter-05-formatted.docx --report verification-report.md
        """
    )

    parser.add_argument('document', help='Path to formatted Word document (.docx)')
    parser.add_argument('--report', help='Output verification report file (default: print to stdout)')

    args = parser.parse_args()

    # Validate inputs
    doc_path = Path(args.document)
    if not doc_path.exists():
        print(f"Error: Document file not found: {doc_path}")
        sys.exit(1)

    if not doc_path.suffix.lower() == '.docx':
        print(f"Error: Document must be .docx format, got: {doc_path.suffix}")
        sys.exit(1)

    # Load document
    print("PacktPub Document Verification")
    print("=" * 70)
    print(f"Document: {doc_path}")
    print()

    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error: Cannot open document: {e}")
        sys.exit(1)

    # Run verifications
    print("Running verifications...")

    style_results = verify_packt_styles(doc)
    print(f"  Styles: {style_results.summary()}")

    code_results = verify_code_blocks(doc)
    print(f"  Code blocks: {code_results.summary()}")

    list_results = verify_list_styles(doc)
    print(f"  Lists: {list_results.summary()}")

    image_results = verify_images(doc)
    print(f"  Images: {image_results.summary()}")

    heading_results = verify_heading_structure(doc)
    print(f"  Headings: {heading_results.summary()}")

    # Generate report
    success = generate_report(
        doc_path,
        style_results,
        code_results,
        list_results,
        image_results,
        heading_results,
        args.report
    )

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
