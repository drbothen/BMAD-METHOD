#!/usr/bin/env python3
"""
Apply PacktPub [PACKT] Styles to Pandoc-Converted Word Documents v5

FIXED:
- Correctly distinguishes bullet lists from numbered lists by checking numFmt attribute
- Splits multi-line code blocks into separate paragraphs with Code End [PACKT] on last line

PacktPub Style System:
- Headings use standard "Heading 1-6" (NO [PACKT] suffix)
- All other content uses [PACKT] suffix (Normal [PACKT], Code [PACKT], etc.)
- Lists have numbering properties but Pandoc assigns "Normal" style
- Code blocks: Code [PACKT] for all lines, Code End [PACKT] for last line

Usage:
    python3 apply-packt-styles-v5.py input.docx output.docx
"""

from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
import sys
from pathlib import Path

# Mapping from Pandoc's built-in styles to PacktPub styles
STYLE_MAPPINGS = {
    # Paragraph styles (add [PACKT] suffix)
    'Normal': 'Normal [PACKT]',
    'Source Code': 'Code [PACKT]',
    'Code': 'Code [PACKT]',
    'Block Quote': 'Quote [PACKT]',
    'Quote': 'Quote [PACKT]',

    # List styles (add [PACKT] suffix)
    # Note: Pandoc creates lists with numbering but keeps "Normal" style
    # We detect these by checking for numPr (numbering properties)
    'List Bullet': 'Bullet [PACKT]',
    'List Number': 'Numbered Bullet [PACKT]',
    'List Paragraph': 'Bullet [PACKT]',
}

# Character style mappings (for inline formatting)
CHARACTER_STYLE_MAPPINGS = {
    'Strong': 'Key Word [PACKT]',
    'Emphasis': 'Italics [PACKT]',
    'Verbatim Char': 'Code In Text [PACKT]',
    'Source Text': 'Code In Text [PACKT]',
}

def has_numbering(paragraph):
    """Check if paragraph has numbering properties (is part of a list)"""
    if paragraph._element.pPr is None:
        return False

    numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    return numPr is not None

def get_numbering_level(paragraph):
    """Get the numbering level (0-based) of a list item"""
    if not has_numbering(paragraph):
        return None

    numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    ilvl = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')

    if ilvl is not None:
        return int(ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0))

    return 0

def is_numbered_list(paragraph, doc):
    """
    Check if this is a numbered list (vs bullet list) by examining the numFmt attribute.

    The correct approach:
    1. Get the paragraph's numId from its numPr
    2. Look up that numId in the numbering definitions to get abstractNumId
    3. Look up the abstractNumId to get the numbering format
    4. Check if level 0's numFmt is 'bullet' or 'decimal'/'lowerLetter'/etc.
    """
    if not has_numbering(paragraph):
        return False

    # Get numId from paragraph
    numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    numId_elem = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')

    if numId_elem is None:
        return False

    numId = numId_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')

    # Get numbering part
    numbering_part = doc.part.numbering_part
    if not numbering_part:
        return False

    # Find the num element with this numId to get abstractNumId
    abstractNumId = None
    for num in numbering_part.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'):
        if num.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId') == numId:
            abstractNumId_elem = num.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
            if abstractNumId_elem is not None:
                abstractNumId = abstractNumId_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            break

    if abstractNumId is None:
        return False

    # Find the abstractNum element with this abstractNumId
    for abstractNum in numbering_part.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum'):
        if abstractNum.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId') == abstractNumId:
            # Get the level 0 numFmt
            for lvl in abstractNum.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lvl'):
                ilvl = lvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl', '0')
                if ilvl == '0':
                    numFmt = lvl.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt')
                    if numFmt is not None:
                        numFmt_val = numFmt.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        # If numFmt is 'bullet', it's a bullet list; anything else is numbered
                        return numFmt_val != 'bullet'
            break

    # Default to bullet if we can't determine
    return False

def split_code_block(paragraph, doc, code_style, code_end_style):
    """
    Split a multi-line code block paragraph into multiple paragraphs.
    Returns a list of new paragraph elements to insert.
    """
    text = paragraph.text
    lines = text.split('\n')

    # If single line, no split needed
    if len(lines) <= 1:
        return None

    # Create new paragraph elements for each line
    new_paras = []

    for i, line in enumerate(lines):
        # Create new paragraph element
        p = OxmlElement('w:p')

        # Create paragraph properties
        pPr = OxmlElement('w:pPr')

        # Add style
        pStyle = OxmlElement('w:pStyle')
        if i == len(lines) - 1:
            # Last line uses Code End [PACKT]
            pStyle.set(qn('w:val'), code_end_style.style_id)
        else:
            # Other lines use Code [PACKT]
            pStyle.set(qn('w:val'), code_style.style_id)
        pPr.append(pStyle)
        p.append(pPr)

        # Add text run
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r.append(t)
        p.append(r)

        new_paras.append(p)

    return new_paras

def apply_packt_styles(input_file, output_file):
    """
    Apply PacktPub [PACKT] styles to a Word document.

    Args:
        input_file: Path to input .docx file (Pandoc output)
        output_file: Path to output .docx file (with PACKT styles)
    """
    print(f"Loading document: {input_file}")
    doc = Document(input_file)

    # Check if PACKT styles exist in the document
    available_styles = {style.name for style in doc.styles}
    packt_styles = {s for s in available_styles if '[PACKT]' in s}

    if not packt_styles:
        print("⚠️  WARNING: No [PACKT] styles found in document!")
        print("   Make sure you used Sample Chapter.docx as --reference-doc")
        return False

    print(f"✓ Found {len(packt_styles)} [PACKT] styles in document")

    # Get style objects
    code_style = doc.styles['Code [PACKT]']
    code_end_style = doc.styles['Code End [PACKT]']

    # First pass: identify code blocks to split
    code_blocks_to_split = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name in ['Source Code', 'Code']:
            if '\n' in para.text:
                code_blocks_to_split.append(i)

    # Second pass: split code blocks (work backwards to preserve indices)
    code_splits = 0
    for idx in reversed(code_blocks_to_split):
        para = doc.paragraphs[idx]
        new_paras = split_code_block(para, doc, code_style, code_end_style)

        if new_paras:
            # Insert new paragraphs after current one
            para_element = para._element
            parent = para_element.getparent()
            para_index = parent.index(para_element)

            # Insert all new paragraphs in order
            for i, new_para in enumerate(new_paras):
                parent.insert(para_index + 1 + i, new_para)

            # Remove original paragraph
            parent.remove(para_element)
            code_splits += 1

    if code_splits > 0:
        print(f"✓ Split {code_splits} multi-line code blocks")

    # Third pass: apply paragraph style mappings
    para_count = 0
    para_mapped = 0
    para_skipped_headings = 0
    lists_mapped = 0
    bullets_mapped = 0
    numbered_mapped = 0
    code_mapped = 0

    for para in doc.paragraphs:
        para_count += 1
        current_style = para.style.name

        # Skip headings - PacktPub uses standard "Heading 1-6" without [PACKT]
        if current_style.startswith('Heading'):
            para_skipped_headings += 1
            continue

        # Skip code blocks - already handled
        if current_style in ['Code [PACKT]', 'Code End [PACKT]']:
            code_mapped += 1
            continue

        # Check if this is a list item (has numbering)
        if has_numbering(para):
            # Determine if bullet or numbered list
            if is_numbered_list(para, doc):
                target_style = 'Numbered Bullet [PACKT]'
                numbered_mapped += 1
            else:
                target_style = 'Bullet [PACKT]'
                bullets_mapped += 1

            if target_style in available_styles:
                para.style = target_style
                lists_mapped += 1
            else:
                print(f"⚠️  List style not found: {target_style}")
            continue

        # Regular paragraph style mapping
        if current_style in STYLE_MAPPINGS:
            target_style = STYLE_MAPPINGS[current_style]

            # Check if target style exists
            if target_style in available_styles:
                para.style = target_style
                para_mapped += 1
            else:
                print(f"⚠️  Style not found: {target_style}")

    print(f"\n✓ Processed {para_count} paragraphs")
    print(f"✓ Mapped {para_mapped} paragraph styles to [PACKT] equivalents")
    print(f"✓ Mapped {code_mapped} code lines (with Code/Code End [PACKT])")
    print(f"✓ Mapped {lists_mapped} list items: {bullets_mapped} bullet, {numbered_mapped} numbered")
    print(f"✓ Kept {para_skipped_headings} headings as standard \"Heading X\" (correct for PacktPub)")

    # Apply character style mappings (runs within paragraphs)
    run_count = 0
    run_mapped = 0

    for para in doc.paragraphs:
        for run in para.runs:
            run_count += 1
            if run.style and run.style.name in CHARACTER_STYLE_MAPPINGS:
                target_style = CHARACTER_STYLE_MAPPINGS[run.style.name]
                if target_style in available_styles:
                    run.style = target_style
                    run_mapped += 1

    if run_mapped > 0:
        print(f"✓ Mapped {run_mapped} / {run_count} character styles to [PACKT] equivalents")

    # Save output
    print(f"\nSaving to: {output_file}")
    doc.save(output_file)
    print("✅ Complete!")

    return True

def print_style_report(input_file):
    """Print a report of styles used in the document"""
    doc = Document(input_file)

    style_usage = {}
    for para in doc.paragraphs:
        style_name = para.style.name
        style_usage[style_name] = style_usage.get(style_name, 0) + 1

    print("\n" + "=" * 70)
    print("STYLE USAGE REPORT")
    print("=" * 70)

    for style, count in sorted(style_usage.items(), key=lambda x: -x[1]):
        if style.startswith('Heading'):
            marker = ' ✓ (PacktPub standard)'
        elif '[PACKT]' in style:
            marker = ' ✓ [PACKT]'
        else:
            marker = ' ⚠️  (needs mapping)'
        print(f"  {count:3}x  {style:45} {marker}")

    packt_count = sum(1 for s in style_usage if '[PACKT]' in s)
    heading_count = sum(1 for s in style_usage if s.startswith('Heading'))
    total_count = len(style_usage)
    correct_count = packt_count + heading_count

    print(f"\nPacktPub-ready styles: {correct_count} / {total_count} style types")
    print(f"  • [PACKT] styles: {packt_count}")
    print(f"  • Standard headings: {heading_count}")
    print("=" * 70 + "\n")

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 apply-packt-styles-v5.py input.docx output.docx")
        print("\nApplies PacktPub [PACKT] styles to Pandoc-converted documents")
        print("\nPacktPub Style System:")
        print("  • Headings: Standard 'Heading 1-6' (no [PACKT])")
        print("  • Lists: Bullet [PACKT] / Numbered Bullet [PACKT]")
        print("  • Code blocks: Code [PACKT] / Code End [PACKT] (last line)")
        print("  • Other content: Add [PACKT] suffix")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2])

    if not input_file.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    print("PacktPub Style Application Tool v5")
    print("=" * 70 + "\n")

    # Apply styles
    success = apply_packt_styles(input_file, output_file)

    if success:
        # Print report
        print_style_report(output_file)
    else:
        sys.exit(1)

if __name__ == "__main__":
    main()
